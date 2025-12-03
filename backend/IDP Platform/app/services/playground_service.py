from fastapi import UploadFile, HTTPException
from bson import ObjectId
import json
import re
import io
import base64

import fitz  # PyMuPDF
from docx import Document
from PIL import Image
from pdf2image import convert_from_bytes

import httpx
import requests
import os

from app.services.qdrant_service import QdrantVector
from app.services.chroma_service import ChromaVector
from app.services.faiss_service import FaissVector
from app.core.config import OPENROUTER_API_KEY, OPENROUTER_MODEL
from app.routers.projects import PROJECT_BASE_PATH

# Separate model for OCR (Gemma-3 vision model)
OCR_MODEL_NAME = "google/gemma-3-12b-it:free"
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"


class PlaygroundService:

    def __init__(self, fs_bucket, projects_collection, extractions_collection):
        self.fs = fs_bucket
        self.projects = projects_collection
        self.extractions = extractions_collection

    async def _get_project_folder(self, project_id: str):
    # Fetch project from DB
        project = await self.projects.find_one({"_id": ObjectId(project_id)})
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")

        project_name = project.get("project_name")
        if not project_name:
            raise HTTPException(status_code=500, detail="project_name missing in DB")

        # Construct folder path (no DB storage required)
        project_folder = os.path.join(PROJECT_BASE_PATH, project_name)
        # Ensure folder exists
        if not os.path.exists(project_folder):
            raise HTTPException(
                status_code=500,
                detail=f"Project folder does not exist: {project_folder}"
            )

        return project_folder

    async def upload_file(self, project_id: str, file: UploadFile):
        # 1) Read file bytes
        file_bytes = await file.read()

        # 2) Store in GridFS
        grid_in = await self.fs.upload_from_stream(
            file.filename,
            file_bytes,
            metadata={
                "project_id": project_id,
                "content_type": file.content_type,
                "size": len(file_bytes),
            }
        )

        # 3) Get project’s local folder
        project_folder = await self._get_project_folder(project_id)

        # 4) Save locally
        local_path = os.path.join(project_folder, file.filename)
        try:
            with open(local_path, "wb") as f:
                f.write(file_bytes)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Saving file failed: {str(e)}")

        # 5) Return success result
        return {
            "status": "success",
            "file_id": str(grid_in),
            "filename": file.filename,
            "saved_path": local_path.replace("\\", "/"),
            "content_type": file.content_type,
            "size": len(file_bytes)
        }

    async def _read_file_from_gridfs(self, file_id: str):
        try:
            grid_out = await self.fs.open_download_stream(ObjectId(file_id))
        except Exception:
            raise HTTPException(status_code=404, detail="File not found in GridFS")

        data = await grid_out.read()
        filename = grid_out.filename
        content_type = grid_out.metadata.get("content_type") if grid_out.metadata else None

        return data, filename, content_type

    async def download_file(self, file_id: str):
        try:
            grid_out = await self.fs.open_download_stream(ObjectId(file_id))
        except Exception:
            raise HTTPException(status_code=404, detail="File not found")

        data = await grid_out.read()
        return grid_out, data
    def _extract_from_pdf(self, data: bytes):
        """
        For searchable PDFs: extract text + tables using PyMuPDF.
        """
        doc = fitz.open(stream=data, filetype="pdf")

        all_text = ""
        tables = []

        for page in doc:
            all_text += page.get_text()

            page_tables = page.find_tables()
            if page_tables:
                for table in page_tables.tables:
                    tables.append(table.extract())

        return all_text, tables

    def _pil_to_data_url(self, img: Image.Image, fmt: str = "PNG") -> str:
        """Convert a PIL Image to base64 data URL."""
        buffer = io.BytesIO()
        img.save(buffer, format=fmt)
        b64 = base64.b64encode(buffer.getvalue()).decode("utf-8")
        return f"data:image/{fmt.lower()};base64,{b64}"

    def _call_gemma_image_ocr(self, image_data_url: str) -> str:
        """
        Call Gemma-3 (via OpenRouter) with one image (data URL) and get structured text.
        """
        if not OPENROUTER_API_KEY:
            raise HTTPException(
                status_code=500,
                detail="OPENROUTER_API_KEY is not configured on the server.",
            )

        headers = {
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json",
        }

        system_prompt = (
            "You are a document OCR assistant. Convert the input image into structured text.\n"
            "Rules:\n"
            "1. Preserve layout as much as possible:\n"
            "   - Keep section headings and subheadings\n"
            "   - Maintain bullet/numbered lists\n"
            "   - Represent tables as Markdown tables\n"
            "   - Keep label: value pairs on the same line.\n"
            "2. Do not summarize or invent content.\n"
            "3. If some text is unreadable, write [UNREADABLE]."
        )

        user_prompt = (
            "Convert this document page image into structured text while preserving headings, "
            "tables (as markdown), and lists."
        )

        body = {
            "model": OCR_MODEL_NAME,
            "messages": [
                {"role": "system", "content": system_prompt},
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": user_prompt},
                        {"type": "image_url", "image_url": {"url": image_data_url}},
                    ],
                },
            ],
            "temperature": 0.1,
            "max_tokens": 4096,
        }

        resp = requests.post(OPENROUTER_URL, headers=headers, json=body)
        try:
            resp.raise_for_status()
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"OCR model request failed: {str(e)} | {resp.text}",
            )

        data = resp.json()
        return data["choices"][0]["message"]["content"]

    def _ocr_extract_pdf(self, data: bytes) -> str:
        """
        OCR for scanned PDFs:
        - Convert each page to image using pdf2image
        - Run Gemma-3 OCR per page
        - Concatenate with page separators
        """
        pages = convert_from_bytes(data, dpi=200)
        all_page_texts = []

        for page_index, page_img in enumerate(pages, start=1):
            data_url = self._pil_to_data_url(page_img, fmt="PNG")
            page_text = self._call_gemma_image_ocr(data_url)
            page_block = f"===== PAGE {page_index} =====\n\n{page_text}"
            all_page_texts.append(page_block)

        return "\n\n\n".join(all_page_texts)

    def _ocr_extract_image(self, data: bytes) -> str:
        """OCR for a single image file (JPG/PNG/etc.) using Gemma."""
        img = Image.open(io.BytesIO(data))
        data_url = self._pil_to_data_url(img)
        return self._call_gemma_image_ocr(data_url)

    def _extract_text(self, data: bytes, filename: str, content_type: str):
        """
        Main text extraction router:
        - Images        → Gemma OCR
        - PDFs scanned  → Gemma OCR (per page)
        - PDFs text     → PyMuPDF (_extract_from_pdf)
        - DOCX          → python-docx
        - TXT           → raw bytes decode
        """
        content_type = content_type or ""
        filename = filename or ""

        filename_lower = filename.lower()

        # 1) Images → Gemma OCR
        if content_type.startswith("image/"):
            ocr_text = self._ocr_extract_image(data)
            return ocr_text, []

        # 2) PDFs
        if filename_lower.endswith(".pdf"):
            # First check if it's searchable or scanned
            doc = fitz.open(stream=data, filetype="pdf")
            has_text = False

            for page in doc:
                page_text = page.get_text().strip()
                if page_text:
                    has_text = True
                    break

            if has_text:
                # Use standard text+table extraction for searchable PDFs
                return self._extract_from_pdf(data)
            else:
                # Scanned PDF → OCR with Gemma
                ocr_text = self._ocr_extract_pdf(data)
                # Tables are embedded as markdown in text
                return ocr_text, []

        # 3) DOCX
        if filename_lower.endswith(".docx"):
            f = io.BytesIO(data)
            doc = Document(f)
            content = "\n".join(p.text for p in doc.paragraphs)
            return content, []

        # 4) TXT
        if filename_lower.endswith(".txt"):
            return data.decode("utf-8", errors="ignore"), []

        # Unknown type
        return "", []

    def _clean_llm_json(self, text: str):
        """
        Cleans LLM output by removing ```json codeblocks,
        stripping markdown and extracting valid JSON.
        Returns a Python dict.
        """
        if not text:
            return {}

        # Remove markdown code block markers
        text = re.sub(r"```json", "", text, flags=re.IGNORECASE).strip()
        text = re.sub(r"```", "", text).strip()

        # Attempt direct JSON load
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            pass

        # Try to extract only the {...} portion
        match = re.search(r"{.*}", text, re.DOTALL)
        if match:
            try:
                return json.loads(match.group(0))
            except json.JSONDecodeError:
                pass

        # If everything fails → throw clear error
        raise HTTPException(
            status_code=500,
            detail="LLM returned invalid JSON; unable to parse.",
        )

    def _build_schema_and_prompt(self, extraction_schema: dict):
        target_schema = {}
        field_prompts = ""

        for field, info in extraction_schema.items():
            target_schema[field] = info.get("type", "string")
            field_prompts += f"- {field}: {info.get('prompt', '')}\n"

        return target_schema, field_prompts
    async def _run_llm_openrouter(
        self,
        project_prompt: str,
        field_prompts: str,
        schema: dict,
        extracted_text: str,
        tables: list,
    ):
        final_prompt = f"""
You are a structured document extraction model.

--- DOCUMENT TEXT ---
{extracted_text}

--- TABLES FOUND ---
{tables}

--- EXTRACTION INSTRUCTIONS ---
{project_prompt}

--- FIELD-SPECIFIC DETAILS ---
{field_prompts}

--- STRICT JSON SCHEMA ---
{schema}

Rules:
1. Output strictly valid JSON.
2. Include ALL fields in the schema.
3. Do NOT include explanations.
        """

        url = "https://openrouter.ai/api/v1/chat/completions"

        headers = {
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "HTTP-Referer": "http://localhost",
            "Content-Type": "application/json",
        }

        payload = {
            "model": OPENROUTER_MODEL,
            "max_tokens": 2000,
            "messages": [
                {"role": "system", "content": "You extract structured data from documents."},
                {"role": "user", "content": final_prompt},
            ],
        }

        async with httpx.AsyncClient(
            timeout=httpx.Timeout(120.0, connect=20.0)
        ) as client:
            response = await client.post(url, json=payload, headers=headers)

        if response.status_code != 200:
            raise HTTPException(status_code=500, detail=response.text)

        data = response.json()
        return data["choices"][0]["message"]["content"]
    async def run_extraction(self, project_id: str, document_type: str, file_id: str):
        # 1) Read file from GridFS
        data, filename, content_type = await self._read_file_from_gridfs(file_id)

        # 2) Extract plain text (via PyMuPDF / DOCX / TXT / Gemma OCR) + tables
        extracted_text, tables = self._extract_text(data, filename, content_type)

        # 3) Get project config
        project = await self.projects.find_one({"_id": ObjectId(project_id)})
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")

        project_prompt = project.get("domain_template", "") or project.get("prompt", "")
        extraction_schema = project.get("extraction_schema", {}) or {}

        # 4) Build target schema + per-field prompts
        target_schema, field_prompts = self._build_schema_and_prompt(extraction_schema)

        # 5) Call OpenRouter LLM for structured extraction
        llm_output = await self._run_llm_openrouter(
            project_prompt,
            field_prompts,
            target_schema,
            extracted_text,
            tables,
        )

        # 6) Clean and parse JSON
        clean_json = self._clean_llm_json(llm_output)

        # 7) Persist extraction result
        await self.extractions.insert_one(
            {
                "project_id": project_id,
                "file_id": file_id,
                "document_type": document_type,
                "schema_used": target_schema,
                "tables": tables,
                "result": clean_json,
            }
        )

        # 8) Return response
        return {
            "status": "success",
            "extracted_data": clean_json,
            "schema_used": target_schema,
            "tables": tables,
        }

    async def get_history_by_project(self, project_id: str):
        return await self.extractions.find({"project_id": project_id}).to_list(1000)

    async def get_history_by_file(self, file_id: str):
        return await self.extractions.find({"file_id": file_id}).to_list(1000)
    
    async def get_extraction_details(self, file_id: str):

        # Fetch record by file_id
        record = await self.extractions.find_one({"file_id": file_id})

        if not record:
            raise HTTPException(status_code=404, detail="Extraction result not found")

        # Extract only what you need
        extraction_result = {
            "status": "success",    # Since your DB does not store status
            "extracted_data": record.get("result")  # <-- Correct field
        }

        return {
            #"file_id": file_id,
            "extraction_result": extraction_result
        }
